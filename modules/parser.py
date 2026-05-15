"""
Module 2 — File Handling & Parsing
Validates uploads, extracts clean text from PDF and DOCX files.
"""

import os
import logging
import pdfplumber
from pypdf import PdfReader
from docx import Document
from werkzeug.utils import secure_filename
from config import Config

logger = logging.getLogger(__name__)


# ──────────────────────────────────────────────
#  Validation Helpers
# ──────────────────────────────────────────────

def allowed_file(filename: str) -> bool:
    """Check that the file has an allowed extension."""
    return (
        "." in filename
        and filename.rsplit(".", 1)[1].lower() in Config.ALLOWED_EXTENSIONS
    )


def safe_save(file_storage) -> tuple[str, str]:
    """
    Validate and save an uploaded FileStorage object.
    Returns (filepath, original_filename) or raises ValueError.
    """
    filename = file_storage.filename
    if not filename:
        raise ValueError("No file selected.")

    if not allowed_file(filename):
        raise ValueError(
            f"Unsupported file type. Please upload a PDF or DOCX file."
        )

    safe_name = secure_filename(filename)
    # Ensure folder exists just-in-time for Vercel /tmp
    if not os.path.exists(Config.UPLOAD_FOLDER):
        os.makedirs(Config.UPLOAD_FOLDER, exist_ok=True)
    
    filepath   = os.path.join(Config.UPLOAD_FOLDER, safe_name)
    file_storage.save(filepath)

    # Enforce size limit after saving (Flask MAX_CONTENT_LENGTH catches most,
    # but this is a belt-and-suspenders check)
    size_mb = os.path.getsize(filepath) / (1024 * 1024)
    if size_mb > Config.MAX_FILE_SIZE_MB:
        os.remove(filepath)
        raise ValueError(
            f"File too large ({size_mb:.1f} MB). Maximum allowed size is "
            f"{Config.MAX_FILE_SIZE_MB} MB."
        )

    return filepath, safe_name


# ──────────────────────────────────────────────
#  Text Extraction
# ──────────────────────────────────────────────

def _extract_pdf(filepath: str) -> str:
    """
    Try pdfplumber first (layout-aware), fall back to pypdf.
    Raises ValueError if the PDF is image-based / empty.
    """
    text = ""

    # Attempt 1: pdfplumber — better for columns & tables
    try:
        with pdfplumber.open(filepath) as pdf:
            pages_text = [page.extract_text() or "" for page in pdf.pages]
            text = "\n".join(pages_text).strip()
    except Exception as e:
        logger.warning("pdfplumber failed (%s), trying pypdf fallback.", e)

    # Attempt 2: pypdf fallback
    if not text:
        try:
            reader = PdfReader(filepath)
            text = "\n".join(
                page.extract_text() or "" for page in reader.pages
            ).strip()
        except Exception as e:
            logger.error("pypdf also failed: %s", e)

    if not text:
        raise ValueError(
            "Your resume appears to be a scanned image or an empty PDF. "
            "Please upload a text-based PDF or a DOCX file."
        )

    return text


def _extract_docx(filepath: str) -> str:
    """Extract text from a DOCX file, including table cells."""
    doc = Document(filepath)
    paragraphs = [p.text for p in doc.paragraphs]

    # Also pull text from tables (common in resume templates)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                paragraphs.append(cell.text)

    text = "\n".join(paragraphs).strip()
    if not text:
        raise ValueError("The uploaded DOCX file appears to be empty.")
    return text


def extract_text(filepath: str) -> str:
    """
    Main entry point: detect format and extract text.
    Always cleans up the file via try/finally.
    """
    try:
        ext = filepath.rsplit(".", 1)[1].lower()
        if ext == "pdf":
            return _extract_pdf(filepath)
        elif ext == "docx":
            return _extract_docx(filepath)
        else:
            raise ValueError("Unsupported file format.")
    finally:
        # Guaranteed cleanup — file is deleted whether or not parsing succeeded
        if os.path.exists(filepath):
            os.remove(filepath)
            logger.debug("Cleaned up temp file: %s", filepath)
